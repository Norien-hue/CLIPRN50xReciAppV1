import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_PATH = r"C:\Users\FA506NC\Desktop\Subida de nota\Interfaces_PI\Documentacion.docx"

DARK_GREEN = RGBColor(0x2E, 0x7D, 0x32)
MED_GREEN = RGBColor(0x38, 0x8E, 0x3C)
MED_GREEN_HEX = "388E3C"
BLACK = RGBColor(0, 0, 0)

doc = Document()

# -- page setup --
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

# -- default style --
sty = doc.styles['Normal']
sty.font.name = 'Calibri'
sty.font.size = Pt(11)
sty.font.color.rgb = BLACK

# ==== helpers ====
def page_number(sec):
    f = sec.footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld)
    run2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run2._r.append(instr)
    run3 = p.add_run()
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld2)

def green_border(p, color=MED_GREEN_HEX, sz="6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    sz = Pt(14) if level == 1 else Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = sz
    r.font.color.rgb = DARK_GREEN
    r.font.name = 'Calibri'
    green_border(p)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(13)
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    return p

# =============================================
# PAGE 1 — PORTADA
# =============================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Universidad de Alicante")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = DARK_GREEN
r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Interfaces de Usuario")
r.font.size = Pt(14)
r.font.color.rgb = MED_GREEN
r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Reciclaje Inteligente\ncon Visión por Computador y CLIP")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = DARK_GREEN
r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Aplicación JavaFX con identificación de productos\nmediante CLIP ViT-B/32")
r.font.size = Pt(12)
r.font.color.rgb = MED_GREEN
r.font.name = 'Calibri'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Alumno: Interfaces PI")
r.font.size = Pt(12)
r.font.color.rgb = BLACK
r.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Curso 2025\u20132026")
r.font.size = Pt(12)
r.font.color.rgb = DARK_GREEN
r.font.name = 'Calibri'

page_number(doc.sections[0])

# =============================================
# PAGE 2 — ÍNDICE
# =============================================
doc.add_page_break()
heading(doc, "\u00cdndice")

toc = [
    ("1. Introducci\u00f3n", 3),
    ("2. Arquitectura del Sistema", 4),
    ("3. C\u00e1mara y Procesamiento de V\u00eddeo", 5),
    ("4. Eliminaci\u00f3n de Fondo", 6),
    ("5. Tabla y Umbral de Confianza", 7),
    ("6. CLIP ViT-B/32 y Coincidencia H\u00edbrida", 8),
    ("7. Integraci\u00f3n con API REST", 9),
    ("8. Despliegue en AWS EC2", 10),
    ("9. Problemas y Soluciones", 11),
    ("10. Conclusiones", 12),
    ("11. Bibliograf\u00eda", 13),
]

for title, pg in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(title)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK_GREEN
    r2 = p.add_run("\t" + str(pg))
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    r2.font.color.rgb = BLACK

# =============================================
# PAGE 3 — 1. INTRODUCCIÓN
# =============================================
doc.add_page_break()
heading(doc, "1. Introducci\u00f3n")

body(doc, (
    "El presente documento describe el desarrollo de un sistema de reciclaje inteligente "
    "realizado en el marco de la asignatura Interfaces de Usuario de la Universidad de "
    "Alicante. El proyecto aborda la problem\u00e1tica social de la clasificaci\u00f3n de residuos "
    "mediante una aplicaci\u00f3n de escritorio con interfaz gr\u00e1fica basada en JavaFX, integrada "
    "con un modelo de inteligencia artificial CLIP (Contrastive Language\u2013Image Pre-training) "
    "para la identificaci\u00f3n autom\u00e1tica de productos a trav\u00e9s de la c\u00e1mara del dispositivo."
))

body(doc, (
    "Uno de los principales problemas que enfrentan los ciudadanos en su d\u00eda a d\u00eda es la "
    "incertidumbre sobre c\u00f3mo clasificar correctamente los residuos. La incorrecta separaci\u00f3n "
    "de materiales reciclables, org\u00e1nicos y no reciclables genera un impacto negativo en el "
    "medio ambiente y reduce la eficiencia de las plantas de tratamiento. Para dar respuesta "
    "a esta necesidad, se ha dise\u00f1ado un sistema automatizado que permite al usuario capturar "
    "una imagen del producto, procesarla mediante t\u00e9cnicas de visi\u00f3n por computador y obtener "
    "una clasificaci\u00f3n precisa del tipo de residuo y su contenedor correspondiente."
))

body(doc, (
    "La soluci\u00f3n consta de una aplicaci\u00f3n JavaFX que gestiona la interacci\u00f3n con el usuario, "
    "el acceso a la c\u00e1mara y la comunicaci\u00f3n con servicios remotos. Un backend Spring Boot "
    "desplegado en AWS EC2 act\u00faa como intermediario entre la aplicaci\u00f3n cliente y un servicio "
    "especializado que ejecuta el modelo CLIP ViT-B/32 en un contenedor Docker. Esta "
    "arquitectura de tres capas garantiza la separaci\u00f3n de responsabilidades y permite "
    "escalar cada componente de forma independiente."
))

# =============================================
# PAGE 4 — 2. ARQUITECTURA DEL SISTEMA
# =============================================
doc.add_page_break()
heading(doc, "2. Arquitectura del Sistema")

body(doc, (
    "La aplicaci\u00f3n sigue una arquitectura de tres capas bien diferenciadas: cliente JavaFX, "
    "servidor API REST Spring Boot y servicio de inferencia CLIP. El cliente JavaFX implementa "
    "el patr\u00f3n Modelo-Vista-Controlador (MVC) para mantener una separaci\u00f3n limpia entre la "
    "l\u00f3gica de negocio, la interfaz de usuario y los datos. Las vistas se definen mediante "
    "archivos FXML, los controladores gestionan los eventos de usuario y los modelos representan "
    "las entidades del dominio."
))

body(doc, (
    "El servidor API REST, implementado con Spring Boot y desplegado en una instancia AWS EC2, "
    "expone los endpoints necesarios para la autenticaci\u00f3n de usuarios, la consulta de productos "
    "y el registro de reciclajes. La comunicaci\u00f3n entre el cliente y el servidor se realiza "
    "mediante peticiones HTTP con autenticaci\u00f3n basada en tokens JWT (JSON Web Token). El flujo "
    "de autenticaci\u00f3n comienza con el login del usuario mediante credenciales o identificaci\u00f3n "
    "por TAP, tras lo cual el servidor devuelve un token JWT que debe incluirse en el "
    "encabezado Authorization de las peticiones posteriores."
))

body(doc, (
    "El servicio de inferencia CLIP se ejecuta en un contenedor Docker sobre otra instancia "
    "EC2, exponiendo su API en el puerto 8080. Cuando el cliente captura una imagen, esta se "
    "env\u00eda al servidor Spring Boot (puerto 3000), que a su vez la reenv\u00eda al servicio CLIP "
    "para su procesamiento. El modelo CLIP ViT-B/32 genera embeddings tanto de la imagen como "
    "de las descripciones textuales de los productos, calculando la similitud coseno para "
    "determinar la coincidencia m\u00e1s probable."
))

body(doc, (
    "El diagrama de despliegue mostrar\u00eda tres nodos principales: el ordenador del usuario "
    "ejecutando la aplicaci\u00f3n JavaFX, la instancia EC2 t3.small que aloja el servidor Spring "
    "Boot, y una segunda instancia EC2 con Docker ejecutando el contenedor CLIP. Las flechas "
    "indicar\u00edan las conexiones HTTP entre los componentes: el cliente JavaFX se comunica con "
    "Spring Boot en el puerto 3000, y este a su vez con el servicio CLIP en el puerto 8080. "
    "Los tokens JWT fluyen desde el servidor al cliente tras la autenticaci\u00f3n y se reenv\u00edan "
    "en cada petici\u00f3n posterior."
))

# =============================================
# PAGE 5 — 3. CÁMARA Y PROCESAMIENTO DE VIDEO
# =============================================
doc.add_page_break()
heading(doc, "3. C\u00e1mara y Procesamiento de V\u00eddeo")

body(doc, (
    "El m\u00f3dulo de captura de v\u00eddeo utiliza la biblioteca OpenCV a trav\u00e9s de su envoltura "
    "JavaCV, empleando espec\u00edficamente la clase FrameGrabber para acceder a la c\u00e1mara web del "
    "dispositivo. La resoluci\u00f3n de captura se establece en 640\u00d7480 p\u00edxeles, aplic\u00e1ndose "
    "posteriormente un recorte cuadrado centrado de 480\u00d7480 p\u00edxeles para eliminar las zonas "
    "perif\u00e9ricas de la imagen y centrar la atenci\u00f3n en el producto."
))

body(doc, (
    "Uno de los desaf\u00edos principales fue evitar el bloqueo de la interfaz de usuario durante "
    "la inicializaci\u00f3n de la c\u00e1mara, ya que el proceso de apertura del dispositivo puede ser "
    "lento y variable. Para resolverlo, se implement\u00f3 un sistema de arranque as\u00edncrono que "
    "realiza la inicializaci\u00f3n en un hilo secundario mientras muestra una barra de progreso "
    "(ProgressBar) en la interfaz. De esta forma, el usuario recibe retroalimentaci\u00f3n visual "
    "del estado de carga sin que la aplicaci\u00f3n se congele."
))

body(doc, (
    "Adicionalmente, se incorpor\u00f3 un mecanismo de guardia mediante AtomicBoolean para prevenir "
    "la doble inicializaci\u00f3n de la c\u00e1mara en caso de que el usuario hiciera clic m\u00faltiples "
    "veces en el bot\u00f3n de activaci\u00f3n. Esta variable at\u00f3mica garantiza que solo un hilo pueda "
    "iniciar la c\u00e1mara, evitando condiciones de carrera y posibles errores de concurrencia. "
    "Una vez que la c\u00e1mara est\u00e1 operativa, los frames capturados se procesan en tiempo real y "
    "se muestran en un ImageView de JavaFX."
))

# =============================================
# PAGE 6 — 4. ELIMINACIÓN DE FONDO
# =============================================
doc.add_page_break()
heading(doc, "4. Eliminaci\u00f3n de Fondo")

body(doc, (
    "La eliminaci\u00f3n del fondo de las im\u00e1genes capturadas constituy\u00f3 uno de los retos t\u00e9cnicos "
    "m\u00e1s significativos del proyecto. En una primera aproximaci\u00f3n, se implement\u00f3 el algoritmo "
    "GrabCut de OpenCV, que ofrece buenos resultados de segmentaci\u00f3n. Sin embargo, al ejecutarse "
    "sobre la m\u00e1quina virtual de Java (JVM) en Windows, se produc\u00edan fallos cr\u00edticos del tipo "
    "EXCEPTION_ACCESS_VIOLATION, provocando el cierre inesperado de la aplicaci\u00f3n. Estos errores "
    "se deben a que GrabCut es una implementaci\u00f3n nativa que interact\u00faa directamente con la "
    "memoria, y su comportamiento no siempre es compatible con la JVM en todas las plataformas."
))

body(doc, (
    "Como soluci\u00f3n alternativa, se desarroll\u00f3 un algoritmo de relleno por inundaci\u00f3n (flood "
    "fill) sensible a bordes implementado \u00edntegramente en Java puro, eliminando cualquier "
    "dependencia de c\u00f3digo nativo. El algoritmo sigue los siguientes pasos: en primer lugar, "
    "se convierte la imagen a escala de grises; a continuaci\u00f3n, se aplica un operador de "
    "gradiente similar a Sobel para detectar los bordes; posteriormente, se ejecuta un flood "
    "fill desde las cuatro esquinas de la imagen con un umbral configurable; finalmente, se "
    "aplica un filtro de mediana de 3\u00d73 para suavizar los bordes y se pinta el fondo de blanco."
))

body(doc, (
    "Para ofrecer flexibilidad al usuario, se incluy\u00f3 un control deslizante (sensitivity "
    "slider) con valores entre 5 y 100, siendo 30 el valor predeterminado. Este control ajusta "
    "el umbral del flood fill, permitiendo al usuario encontrar el equilibrio \u00f3ptimo entre "
    "eliminaci\u00f3n de fondo y conservaci\u00f3n de detalles del producto. La vista previa se actualiza "
    "en tiempo real, lo que facilita el ajuste interactivo."
))

body(doc, "La siguiente tabla compara ambos enfoques:")

# -- comparison table --
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr = ['Caracter\u00edstica', 'GrabCut (OpenCV)', 'Flood Fill (Java puro)']
rows_data = [
    ['Rendimiento', 'R\u00e1pido (implementaci\u00f3n nativa)', 'M\u00e1s lento (Java puro)'],
    ['Estabilidad', 'Inestable (crashes en Windows)', 'Estable (sin c\u00f3digo nativo)'],
    ['Dependencias', 'OpenCV nativo', 'Ninguna'],
]
for i, h in enumerate(hdr):
    cell = table.rows[0].cells[i]
    cell.text = ''
    r = cell.paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = DARK_GREEN
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
for ri, rd in enumerate(rows_data, 1):
    for ci, val in enumerate(rd):
        cell = table.rows[ri].cells[ci]
        cell.text = ''
        r = cell.paragraphs[0].add_run(val)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r.font.color.rgb = BLACK

# =============================================
# PAGE 7 — 5. TABLA Y UMBRAL DE CONFIANZA
# =============================================
doc.add_page_break()
heading(doc, "5. Tabla y Umbral de Confianza")

body(doc, (
    "Los resultados de la clasificaci\u00f3n de productos se presentan al usuario mediante un "
    "TableView de JavaFX. Cada fila de la tabla muestra el nombre del producto, la puntuaci\u00f3n "
    "de similitud (score) obtenida del modelo CLIP y el tipo de residuo asociado, ordenados "
    "de mayor a menor puntuaci\u00f3n."
))

body(doc, (
    "Se estableci\u00f3 un umbral de confianza del 47 %: aquellos productos cuya puntuaci\u00f3n se "
    "sit\u00fae por debajo de este valor se muestran en color rojo y no pueden ser seleccionados "
    "por el usuario. Esta medida evita que el sistema sugiera productos con baja confianza, "
    "reduciendo la probabilidad de clasificaciones err\u00f3neas. La implementaci\u00f3n utiliza cell "
    "factories personalizadas que aplican estilos condicionales en funci\u00f3n del valor num\u00e9rico "
    "del score."
))

body(doc, (
    "Adem\u00e1s, se incorpor\u00f3 un campo de b\u00fasqueda que permite filtrar los productos por nombre, "
    "facilitando la localizaci\u00f3n de un art\u00edculo concreto cuando el cat\u00e1logo es extenso. La "
    "tabla se actualiza din\u00e1micamente a medida que el usuario escribe en el campo de texto, "
    "aplicando el filtro sobre la lista completa de resultados."
))

# =============================================
# PAGE 8 — 6. CLIP ViT-B/32 Y COINCIDENCIA HÍBRIDA
# =============================================
doc.add_page_break()
heading(doc, "6. CLIP ViT-B/32 y Coincidencia H\u00edbrida")

body(doc, (
    "El n\u00facleo del sistema de identificaci\u00f3n se basa en el modelo CLIP (Contrastive Language\u2013"
    "Image Pre-training) desarrollado por OpenAI, en su variante ViT-B/32. Este modelo emplea "
    "una arquitectura Vision Transformer (ViT) con parches de 32\u00d732 p\u00edxeles para procesar "
    "im\u00e1genes y un transformer para procesar texto, proyectando ambos dominios en un espacio "
    "de embeddings com\u00fan de 512 dimensiones."
))

body(doc, (
    "La estrategia de coincidencia implementada es h\u00edbrida: se combinan los embeddings de la "
    "imagen capturada con los embeddings de las descripciones textuales de los productos, "
    "aplicando un peso de mezcla (BLEND_WEIGHT) de 0.5. Esto significa que la puntuaci\u00f3n final "
    "es el promedio ponderado entre la similitud coseno imagen-imagen y la similitud coseno "
    "texto-imagen. Este enfoque h\u00edbrido mejora la robustez del sistema, ya que no depende "
    "exclusivamente de la calidad visual de la captura."
))

body(doc, (
    "Los embeddings de los ocho productos disponibles en el cat\u00e1logo se precalculan y "
    "almacenan, lo que acelera el proceso de inferencia. La ejecuci\u00f3n del modelo sobre una "
    "instancia t3.small de AWS EC2 tiene un tiempo medio de inferencia de aproximadamente "
    "500 milisegundos y un consumo de memoria RAM cercano a 1.2 GB. La respuesta de depuraci\u00f3n "
    "incluye las puntuaciones individuales de cada producto, lo que facilita el diagn\u00f3stico "
    "y la puesta a punto del sistema."
))

# =============================================
# PAGE 9 — 7. INTEGRACIÓN CON API REST
# =============================================
doc.add_page_break()
heading(doc, "7. Integraci\u00f3n con API REST")

body(doc, (
    "La comunicaci\u00f3n entre la aplicaci\u00f3n JavaFX y el servidor se realiza a trav\u00e9s de una API "
    "RESTful implementada con Spring Boot. El cliente HTTP se implementa mediante "
    "java.net.http.HttpClient, encapsulado en una clase ApiClient que sigue el patr\u00f3n singleton "
    "para garantizar una \u00fanica instancia compartida en toda la aplicaci\u00f3n."
))

body(doc, (
    "Los endpoints disponibles son los siguientes: GET /api/health para la verificaci\u00f3n del "
    "estado del servidor; POST /api/usuarios/login para la autenticaci\u00f3n de usuarios; GET "
    "/api/productos/barcode/{code} para la consulta de productos por c\u00f3digo de barras; GET "
    "/api/usuarios/by-tap/{tap} para la identificaci\u00f3n de usuarios mediante TAP; y POST "
    "/api/historial para el registro de reciclajes, que requiere autenticaci\u00f3n mediante token "
    "JWT."
))

body(doc, (
    "La autenticaci\u00f3n se realiza mediante el esquema Bearer: el token JWT obtenido tras el "
    "login se incluye en el encabezado HTTP Authorization de las peticiones subsiguientes. "
    "Se ha configurado un tiempo de espera de 60 segundos para las peticiones de coincidencia "
    "de productos, tras el cual se muestra un mensaje de error amigable al usuario mediante "
    "alertas de JavaFX (Alert). El manejo de errores contempla tanto los errores de red como "
    "las respuestas HTTP con c\u00f3digos de error, mostrando en cada caso el mensaje adecuado."
))

# =============================================
# PAGE 10 — 8. DESPLIEGUE EN AWS EC2
# =============================================
doc.add_page_break()
heading(doc, "8. Despliegue en AWS EC2")

body(doc, (
    "El despliegue del sistema se realiz\u00f3 sobre instancias EC2 de AWS con sistema operativo "
    "Ubuntu. Se utilizaron dos instancias de tipo t3.small: una para el servidor Spring Boot "
    "y otra para el servicio CLIP en contenedor Docker."
))

body(doc, (
    "El servidor Spring Boot se despliega como un archivo JAR ejecutable. El proceso de "
    "despliegue comienza con la transferencia del archivo JAR mediante scp, seguido de la "
    "conexi\u00f3n SSH a la instancia y la ejecuci\u00f3n con java -jar. La aplicaci\u00f3n queda expuesta "
    "en el puerto 3000."
))

body(doc, (
    "El servicio CLIP se despliega como un contenedor Docker. Se construye una imagen Docker "
    "a partir de un Dockerfile que instala las dependencias necesarias (Python, PyTorch, CLIP) "
    "y expone el puerto 8080. El contenedor se ejecuta con docker run, mapeando el puerto 8080 "
    "del contenedor al puerto 8080 del host. Las variables de entorno se configuran para "
    "parametrizar la conexi\u00f3n a la base de datos y otras opciones del servidor."
))

body(doc, (
    "A continuaci\u00f3n se muestran los comandos principales utilizados para el despliegue: "
    "scp -i clave.pem app.jar ubuntu@ip:/home/ubuntu/ para transferir el JAR; "
    "ssh -i clave.pem ubuntu@ip para conectar a la instancia; "
    "java -jar app.jar --server.port=3000 para iniciar el servidor; "
    "docker build -t clip-service . para construir la imagen CLIP; y "
    "docker run -d -p 8080:8080 clip-service para ejecutar el contenedor."
))

# =============================================
# PAGE 11 — 9. PROBLEMAS Y SOLUCIONES
# =============================================
doc.add_page_break()
heading(doc, "9. Problemas y Soluciones")

heading(doc, "Problema 1: Fallos de GrabCut en Windows", level=2)
body(doc, (
    "El algoritmo GrabCut de OpenCV provocaba errores EXCEPTION_ACCESS_VIOLATION en la JVM "
    "sobre Windows, causando cierres inesperados de la aplicaci\u00f3n. La soluci\u00f3n fue implementar "
    "un algoritmo de flood fill sensible a bordes completamente en Java puro, eliminando las "
    "dependencias de c\u00f3digo nativo."
))

heading(doc, "Problema 2: Congelaci\u00f3n de la interfaz al iniciar la c\u00e1mara", level=2)
body(doc, (
    "La inicializaci\u00f3n de la c\u00e1mara bloqueaba el hilo principal de JavaFX, congelando la "
    "interfaz de usuario. Se resolvi\u00f3 realizando la inicializaci\u00f3n en un hilo secundario y "
    "mostrando una barra de progreso (ProgressBar) para informar al usuario del estado de la "
    "operaci\u00f3n."
))

heading(doc, "Problema 3: Doble inicializaci\u00f3n de la c\u00e1mara", level=2)
body(doc, (
    "Los clics r\u00e1pidos y repetidos en el bot\u00f3n de activaci\u00f3n pod\u00edan iniciar la c\u00e1mara "
    "m\u00faltiples veces, generando condiciones de carrera. La soluci\u00f3n fue emplear un AtomicBoolean "
    "como mecanismo de guardia, permitiendo que solo un hilo pudiera iniciar la c\u00e1mara."
))

heading(doc, "Problema 4: Selecci\u00f3n de productos con baja confianza", level=2)
body(doc, (
    "El sistema pod\u00eda sugerir productos con puntuaciones muy bajas, lo que aumentaba el riesgo "
    "de clasificaciones incorrectas. Se estableci\u00f3 un umbral m\u00ednimo del 47 %: los productos por "
    "debajo de este valor se muestran en rojo y no son seleccionables, utilizando cell factories "
    "personalizadas en el TableView."
))

heading(doc, "Problema 5: Espacio en disco en EC2", level=2)
body(doc, (
    "El disco /var de la instancia EC2 se llenaba debido a im\u00e1genes Docker antiguas y logs "
    "acumulados. La soluci\u00f3n fue ejecutar peri\u00f3dicamente docker system prune -a para eliminar "
    "contenedores, im\u00e1genes y redes no utilizados, liberando espacio en disco."
))

# =============================================
# PAGE 12 — 10. CONCLUSIONES
# =============================================
doc.add_page_break()
heading(doc, "10. Conclusiones")

body(doc, (
    "El sistema de reciclaje inteligente desarrollado demuestra que la combinaci\u00f3n de una "
    "interfaz de usuario JavaFX con un modelo de inteligencia artificial CLIP es una soluci\u00f3n "
    "viable y eficaz para la clasificaci\u00f3n automatizada de residuos. La aplicaci\u00f3n proporciona "
    "una experiencia de usuario fluida y accesible, cumpliendo con los objetivos planteados al "
    "inicio del proyecto."
))

body(doc, (
    "La adopci\u00f3n de un enfoque puramente Java para la eliminaci\u00f3n de fondo, aunque supuso un "
    "esfuerzo de desarrollo adicional, result\u00f3 ser la decisi\u00f3n correcta al evitar los problemas "
    "de estabilidad asociados a las bibliotecas nativas en Windows. El modelo CLIP ViT-B/32 ha "
    "demostrado un rendimiento notable en la identificaci\u00f3n de productos, con tiempos de "
    "inferencia aceptables en infraestructura cloud."
))

body(doc, (
    "Como l\u00edneas de trabajo futuro, se plantea ampliar el cat\u00e1logo de productos reconocibles, "
    "mejorar el algoritmo de eliminaci\u00f3n de fondo mediante t\u00e9cnicas de deep learning, y "
    "desarrollar una versi\u00f3n m\u00f3vil de la aplicaci\u00f3n que permita a los usuarios clasificar "
    "residuos desde sus tel\u00e9fonos inteligentes."
))

# =============================================
# PAGE 13 — 11. BIBLIOGRAFÍA
# =============================================
doc.add_page_break()
heading(doc, "11. Bibliograf\u00eda")

refs = [
    "Oracle Corporation. JavaFX Documentation. https://openjfx.io/",
    "Radford, A., Kim, J. W., Hallacy, C., et al. (2021). Learning Transferable Visual Models From Natural Language Supervision. Proceedings of the International Conference on Machine Learning (ICML).",
    "Bradski, G. (2000). The OpenCV Library. Dr. Dobb\u2019s Journal of Software Tools.",
    "Bytedeco. JavaCV Documentation. https://github.com/bytedeco/javacv",
    "Pivotal Software. Spring Boot Reference Documentation. https://spring.io/projects/spring-boot",
    "OpenAI. CLIP GitHub Repository. https://github.com/openai/CLIP",
    "Docker Inc. Docker Documentation. https://docs.docker.com/",
    "Amazon Web Services. Amazon EC2 Documentation. https://docs.aws.amazon.com/ec2/",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{i}] {ref}")
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK

# ========== SAVE ==========
doc.save(DOCX_PATH)
print("DOCX generated successfully")
